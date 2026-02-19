import logging
import os
import re
import tempfile
import shutil
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
from dotenv import load_dotenv

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import existing pipeline functions
import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from run_batch import (
    extract_text_from_pdf,
    extract_tables_from_pdf,
    clean_text,
    reinsert_tables_into_markdown,
    convert_to_protocol_markdown,
    finalize_protocol_md,
    generate_flags_md,
    append_flags_summary
)

load_dotenv()

app = FastAPI()

@app.get("/health")
def health():
    logger.info("Health check pinged")
    return {"status": "ok", "timestamp": datetime.now().isoformat()}

# CORS: allow Next.js frontend in dev and production
# Set EXTRA_CORS_ORIGINS env var for additional origins (comma-separated)
_origins = [
    "http://localhost:3000",
    "https://protocolfoundry.io",
    "https://www.protocolfoundry.io",
]
_extra = os.getenv("EXTRA_CORS_ORIGINS", "")
if _extra:
    _origins.extend([o.strip() for o in _extra.split(",") if o.strip()])

app.add_middleware(
    CORSMiddleware,
    allow_origins=_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB


def _get_or_create_restart_num_id(doc):
    """Create a new w:num entry with startOverride=1 that shares the same
    abstractNum as the built-in 'List Number' style.  Returns the new numId
    integer, or None if the numbering metadata cannot be located."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        numbering = doc.part.numbering_part._element
    except Exception:
        return None

    # Find the numId referenced by the 'List Number' paragraph style.
    try:
        list_style = doc.styles['List Number']
        style_numPr = list_style.element.find(qn('w:pPr')).find(qn('w:numPr'))
        base_num_id = int(style_numPr.find(qn('w:numId')).get(qn('w:val')))
    except (AttributeError, TypeError, ValueError):
        return None

    # Walk the numbering definitions to find the matching abstractNumId.
    abstract_num_id = None
    for num_elem in numbering.findall(qn('w:num')):
        try:
            if int(num_elem.get(qn('w:numId'))) == base_num_id:
                abstract_num_id = int(
                    num_elem.find(qn('w:abstractNumId')).get(qn('w:val'))
                )
                break
        except (AttributeError, TypeError, ValueError):
            continue

    if abstract_num_id is None:
        return None

    all_ids = [
        int(n.get(qn('w:numId')))
        for n in numbering.findall(qn('w:num'))
        if n.get(qn('w:numId')) is not None
    ]
    new_id = max(all_ids) + 1

    # <w:num w:numId="new_id">
    #   <w:abstractNumId w:val="abstract_num_id"/>
    #   <w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>
    # </w:num>
    new_num = OxmlElement('w:num')
    new_num.set(qn('w:numId'), str(new_id))

    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(abstract_num_id))
    new_num.append(abstract_ref)

    lvl = OxmlElement('w:lvlOverride')
    lvl.set(qn('w:ilvl'), '0')
    so = OxmlElement('w:startOverride')
    so.set(qn('w:val'), '1')
    lvl.append(so)
    new_num.append(lvl)

    numbering.append(new_num)
    return new_id


def _apply_num_id_to_para(para, num_id: int) -> None:
    """Overwrite the numId in a paragraph's numPr so it belongs to the
    specified list sequence rather than the style default."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = para._p.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pPr.append(numPr)

    ilvl = numPr.find(qn('w:ilvl'))
    if ilvl is None:
        ilvl = OxmlElement('w:ilvl')
        numPr.insert(0, ilvl)
    ilvl.set(qn('w:val'), '0')

    num_id_elem = numPr.find(qn('w:numId'))
    if num_id_elem is None:
        num_id_elem = OxmlElement('w:numId')
        numPr.append(num_id_elem)
    num_id_elem.set(qn('w:val'), str(num_id))


def _add_raw_table_to_doc(doc, rows: list) -> None:
    """Render pdfplumber raw table rows (list[list[str|None]]) as a docx Table."""
    if not rows:
        return
    num_cols = max((len(r) for r in rows), default=0)
    if num_cols == 0:
        return
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(num_cols):
            val = row_data[c_idx] if c_idx < len(row_data) else None
            cell_text = str(val).strip().replace('\n', ' ') if val is not None else ''
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = cell_text
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _add_markdown_table_to_doc(doc, table_lines: list) -> None:
    """Create a python-docx Table from markdown table lines."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator row — inner content is only dashes, colons, spaces
        inner = line.strip('|')
        if re.match(r'^[\s\-|:]+$', inner):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx in range(num_cols):
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
            cell = row.cells[c_idx]
            cell.text = cell_text
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _split_table_blocks(table_lines: list) -> list:
    """Split a flat list of consecutive pipe-delimited lines into individual
    table blocks.  When a separator row (cells containing only dashes/colons)
    appears after the first two lines of the current block it signals that the
    line immediately preceding it is the header of a NEW table.  That line is
    moved into a fresh block together with the separator, and the old block is
    closed.  This handles back-to-back tables that share no blank-line boundary.
    """
    if not table_lines:
        return []

    def _is_sep(line: str) -> bool:
        inner = line.strip().strip('|')
        return bool(inner and re.match(r'^[\s\-|:]+$', inner))

    blocks = []
    current = []
    for line in table_lines:
        if _is_sep(line) and len(current) >= 2:
            # Separator found mid-block: everything up to (but not including)
            # the last line belongs to the previous table; the last line becomes
            # the header of the new table.
            blocks.append(current[:-1])
            current = [current[-1], line]
        else:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def markdown_to_docx(md_content: str, output_path: Path, pdf_tables: list = None):
    """Convert markdown to docx using python-docx.

    pdf_tables – ordered list of raw table dicts from extract_tables_from_pdf().
    When present, TABLE_PLACEHOLDER_N lines are replaced with the corresponding
    pre-extracted table rendered directly from pdfplumber row data.
    """
    from docx import Document

    doc = Document()
    lines = md_content.split('\n')
    i = 0
    # Tracks whether the next numbered-list paragraph should restart at 1.
    _restart_next_list = False
    # The numId currently in use for the active section's numbered list.
    _section_num_id = None

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # TABLE_PLACEHOLDER_N — substitute the pre-extracted PDF table
        m_ph = re.match(r'^TABLE_PLACEHOLDER_(\d+)$', line.strip())
        if m_ph:
            idx = int(m_ph.group(1)) - 1  # convert 1-based N to 0-based index
            if pdf_tables and 0 <= idx < len(pdf_tables):
                _add_raw_table_to_doc(doc, pdf_tables[idx]['rows'])
            i += 1
            continue

        # Detect markdown table
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            for block in _split_table_blocks(table_lines):
                _add_markdown_table_to_doc(doc, block)
            continue

        # Handle headings — each heading marks the start of a new list sequence
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            _restart_next_list = True
            _section_num_id = None
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            _restart_next_list = True
            _section_num_id = None
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            _restart_next_list = True
            _section_num_id = None
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            _restart_next_list = True
            _section_num_id = None

        # Handle bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists — match any number of leading digits (1., 10., 12.) etc.
        elif re.match(r'^\d+[.)]\s', line):
            m = re.match(r'^\d+[.)]\s+', line)
            text = line[m.end():]
            para = doc.add_paragraph(text, style='List Number')
            # First numbered item after a heading: create a fresh numbering
            # sequence so this section restarts at 1.
            if _restart_next_list:
                new_id = _get_or_create_restart_num_id(doc)
                if new_id:
                    _section_num_id = new_id
                _restart_next_list = False
            # Pin every item in this section to the same numId so the counter
            # is continuous within the section and isolated from other sections.
            if _section_num_id:
                _apply_num_id_to_para(para, _section_num_id)

        # Regular paragraph
        else:
            doc.add_paragraph(line)

        i += 1

    doc.save(str(output_path))


@app.post("/convert")
async def convert_pdf(file: UploadFile = File(...)):
    """
    Convert uploaded PDF to protocol DOCX.
    """
    # Validate file type
    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    # Read file content
    file_content = await file.read()
    
    # Validate file size
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File size exceeds 20MB limit")

    file_size_mb = len(file_content) / (1024 * 1024)
    logger.info(f"[{datetime.now()}] Conversion started: {file.filename} ({file_size_mb:.2f} MB)")

    # Create temporary directory
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Save uploaded PDF
        pdf_path = temp_dir / "input.pdf"
        with open(pdf_path, 'wb') as f:
            f.write(file_content)
        
        # Initialize OpenAI client
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise HTTPException(status_code=500, detail="OPENAI_API_KEY not configured")
        
        client = OpenAI(api_key=api_key)
        
        # Load contract
        contract_path = Path(__file__).parent.parent.parent / "contract.md"
        if not contract_path.exists():
            raise HTTPException(status_code=500, detail="contract.md not found in repository root")
        
        with open(contract_path, 'r', encoding='utf-8') as f:
            contract = f.read()
        
        # Run pipeline
        try:
            pdf_tables = extract_tables_from_pdf(pdf_path)
            raw_text = extract_text_from_pdf(pdf_path)
            cleaned_text = clean_text(raw_text)
            protocol_md = convert_to_protocol_markdown(client, contract, cleaned_text)
            protocol_md = reinsert_tables_into_markdown(protocol_md, pdf_tables)
            protocol_md = finalize_protocol_md(protocol_md)
            flags_md = generate_flags_md(client, cleaned_text, protocol_md)
            final_md = append_flags_summary(protocol_md, flags_md)
        except Exception as e:
            logger.error(f"[{datetime.now()}] Pipeline failed: {file.filename} - {e}")
            raise HTTPException(status_code=500, detail=f"Pipeline error: {str(e)}")

        # Convert to DOCX — tables are now embedded directly in the markdown
        docx_path = temp_dir / "output.docx"

        try:
            markdown_to_docx(final_md, docx_path, [])
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX conversion failed: {str(e)}")
        
        logger.info(f"[{datetime.now()}] Conversion completed: {file.filename}")

        # Return file
        return FileResponse(
            path=str(docx_path),
            filename="protocol.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    except HTTPException:
        # Clean up temp directory
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise
    
    except Exception as e:
        logger.error(f"[{datetime.now()}] Unexpected error: {file.filename} - {e}")
        # Clean up temp directory
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")