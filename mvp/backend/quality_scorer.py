"""quality_scorer.py

Score reference DOCX files against ground truth JSONs and produce a
timestamped report.

Usage:
    python quality_scorer.py                              # score all 15
    python quality_scorer.py --pdf "MRC5_Cell Factory_Thermo"  # score one
"""

import argparse
import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).parent
DOCX_DIR   = SCRIPT_DIR / "test_data" / "reference_docx"
GT_DIR     = SCRIPT_DIR / "test_data" / "ground_truth"
REPORT_DIR = SCRIPT_DIR / "test_data" / "reports"
REPO_ROOT  = (SCRIPT_DIR / "../..").resolve()

SCORE_HISTORY_PATH = REPORT_DIR / "score_history.json"
VERSION = "1.0"


# ---------------------------------------------------------------------------
# Git helper
# ---------------------------------------------------------------------------

def get_git_commit() -> str:
    try:
        result = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=str(REPO_ROOT),
            capture_output=True,
            text=True,
            timeout=5,
        )
        if result.returncode == 0:
            return result.stdout.strip()
    except Exception:
        pass
    return "unknown"


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def load_docx(docx_path: Path) -> Document:
    import io
    import shutil
    import tempfile
    # Copy to a temp file to avoid OneDrive/path-with-spaces lock issues
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        shutil.copy2(docx_path, tmp.name)
        tmp_path = tmp.name
    try:
        with open(tmp_path, "rb") as f:
            return Document(io.BytesIO(f.read()))
    finally:
        Path(tmp_path).unlink(missing_ok=True)


def get_all_text(doc: Document) -> str:
    """Return all paragraph text joined with newlines."""
    return "\n".join(p.text for p in doc.paragraphs)


def get_heading_and_bold_texts(doc: Document) -> list[str]:
    """Return text of paragraphs that are headings or fully-bold runs."""
    results = []
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if "Heading" in style_name:
            if p.text.strip():
                results.append(p.text.strip())
            continue
        runs = [r for r in p.runs if r.text.strip()]
        if runs and all(r.bold for r in runs):
            results.append(p.text.strip())
    return results


# ---------------------------------------------------------------------------
# Dimension 1: Structural Completeness (30 pts)
# ---------------------------------------------------------------------------

def score_structural(doc: Document, gt: dict) -> tuple[float, dict]:
    expected = gt.get("expected_sections", [])
    if not expected:
        return 30.0, {"sections_expected": 0, "sections_found": 0, "missing": []}

    headings = get_heading_and_bold_texts(doc)
    headings_lower = [h.lower() for h in headings]

    found, missing = [], []
    for section in expected:
        sec_lower = section.lower()
        if any(sec_lower in h for h in headings_lower):
            found.append(section)
        else:
            missing.append(section)

    score = round((len(found) / len(expected)) * 30, 1)
    detail = {
        "sections_expected": len(expected),
        "sections_found": len(found),
        "missing": missing,
    }
    return score, detail


# ---------------------------------------------------------------------------
# Dimension 2: Content Fidelity (40 pts)
# ---------------------------------------------------------------------------

def _extract_numbers(value_str: str) -> list[str]:
    """Extract individual numeric tokens from a value string.

    Handles plain numbers, ranges ('2-3', '8 to 11'), and scientific
    notation ('3 x 10^5').  Returns a list of strings to search for.
    """
    # Remove common exponent/scientific-notation suffixes for plain search
    cleaned = re.sub(r"\s*x\s*10\^\d+", "", value_str)
    # Split on common range separators
    parts = re.split(r"\s*(?:to|-)\s*", cleaned)
    tokens = []
    for part in parts:
        # Extract the first decimal/integer number from each part
        m = re.search(r"\d+(?:[.,]\d+)?", part)
        if m:
            tokens.append(m.group().replace(",", "."))
    return tokens if tokens else [value_str]


def _value_near_hint(full_text: str, number: str, hint: str, window: int = 400) -> bool:
    """Return True if *number* appears as a word boundary within *window*
    characters of *hint* (case-insensitive) in *full_text*."""
    hint_lower = hint.lower()
    text_lower = full_text.lower()
    number_pattern = re.compile(r"\b" + re.escape(number) + r"\b")

    start = 0
    while True:
        hint_pos = text_lower.find(hint_lower, start)
        if hint_pos == -1:
            break
        region_start = max(0, hint_pos - window)
        region_end = min(len(full_text), hint_pos + len(hint_lower) + window)
        region = full_text[region_start:region_end]
        if number_pattern.search(region):
            return True
        start = hint_pos + 1
    return False


def score_content_fidelity(doc: Document, gt: dict) -> tuple[float, dict]:
    critical_values = gt.get("critical_values", [])
    if not critical_values:
        return 40.0, {"values_total": 0, "values_found": 0, "failed": []}

    full_text = get_all_text(doc)
    found_count = 0
    failed = []

    for cv in critical_values:
        cv_id   = cv.get("id", "?")
        desc    = cv.get("description", "")
        value   = cv.get("value", "")
        hint    = cv.get("context_hint", "")
        numbers = _extract_numbers(value)

        passed = any(_value_near_hint(full_text, num, hint) for num in numbers)
        if passed:
            found_count += 1
        else:
            failed.append({"id": cv_id, "description": desc, "value": value})

    score = round((found_count / len(critical_values)) * 40, 1)
    detail = {
        "values_total": len(critical_values),
        "values_found": found_count,
        "failed": failed,
    }
    return score, detail


# ---------------------------------------------------------------------------
# Dimension 3: Formatting Correctness (20 pts)
# ---------------------------------------------------------------------------

def score_formatting(doc: Document, gt: dict) -> tuple[float, dict]:
    checks_expected = gt.get("formatting_checks", {})
    full_text = get_all_text(doc)
    results = {}
    total_pts = 0

    if checks_expected.get("has_numbered_steps"):
        count = sum(
            1 for p in doc.paragraphs
            if (
                re.match(r"^\s*\d+[\.\)]", p.text)
                or "List" in (p.style.name if p.style else "")
                or "Number" in (p.style.name if p.style else "")
            )
        )
        passed = count >= 3
        results["has_numbered_steps"] = {"passed": passed, "count": count, "pts": 5 if passed else 0}
        total_pts += 5 if passed else 0

    if checks_expected.get("headers_use_bold"):
        header_paras = sum(
            1 for p in doc.paragraphs
            if (
                ("Heading" in (p.style.name if p.style else ""))
                or ("Title" in (p.style.name if p.style else ""))
                or (
                    p.runs
                    and all(r.bold for r in p.runs if r.text.strip())
                    and p.text.strip()
                )
            )
        )
        passed = header_paras >= 2
        results["headers_use_bold"] = {"passed": passed, "count": header_paras, "pts": 5 if passed else 0}
        total_pts += 5 if passed else 0

    if checks_expected.get("no_bare_tables"):
        table_count = len(doc.tables)
        passed = table_count == 0
        results["no_bare_tables"] = {"passed": passed, "table_count": table_count, "pts": 5 if passed else 0}
        total_pts += 5 if passed else 0

    if checks_expected.get("inline_value_embedding"):
        count = sum(
            1 for p in doc.paragraphs
            if re.search(r"\(.*\d+.*\)", p.text)
        )
        passed = count >= 3
        results["inline_value_embedding"] = {"passed": passed, "count": count, "pts": 5 if passed else 0}
        total_pts += 5 if passed else 0

    return float(total_pts), results


# ---------------------------------------------------------------------------
# Dimension 4: Hallucination Flag (10 pts)
# ---------------------------------------------------------------------------

def score_hallucination(doc: Document, gt: dict) -> tuple[float, dict]:
    watchlist = gt.get("hallucination_watchlist", [])
    if not watchlist:
        return 10.0, {"triggered": []}

    full_text = get_all_text(doc).lower()
    triggered = [
        term for term in watchlist
        if re.search(re.escape(term.lower()), full_text)
    ]

    score = 0.0 if triggered else 10.0
    return score, {"triggered": triggered}


# ---------------------------------------------------------------------------
# Score one PDF
# ---------------------------------------------------------------------------

def score_pdf(stem: str) -> dict | None:
    docx_path = DOCX_DIR / f"{stem}.docx"
    gt_path   = GT_DIR   / f"{stem}.json"

    if not docx_path.exists():
        return {"stem": stem, "error": f"DOCX not found: {docx_path.name}"}
    if not gt_path.exists():
        return {"stem": stem, "error": f"Ground truth not found: {gt_path.name}"}

    try:
        doc = load_docx(docx_path)
        gt  = json.loads(gt_path.read_text(encoding="utf-8"))

        s1, d1 = score_structural(doc, gt)
        s2, d2 = score_content_fidelity(doc, gt)
        s3, d3 = score_formatting(doc, gt)
        s4, d4 = score_hallucination(doc, gt)

        total = round(s1 + s2 + s3 + s4, 1)
        return {
            "stem":          stem,
            "total":         total,
            "structural":    {"score": s1, "detail": d1},
            "content":       {"score": s2, "detail": d2},
            "formatting":    {"score": s3, "detail": d3},
            "hallucination": {"score": s4, "detail": d4},
        }
    except Exception as exc:
        return {"stem": stem, "error": f"Scoring failed: {exc}"}


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------

def _bar(score: float, max_score: float, width: int = 20) -> str:
    filled = int(round((score / max_score) * width)) if max_score else 0
    return "[" + "#" * filled + "-" * (width - filled) + "]"


def build_report(results: list[dict], timestamp: str, commit: str) -> str:
    scored   = [r for r in results if "error" not in r]
    errored  = [r for r in results if "error" in r]
    lines    = []

    # Header
    lines += [
        "=" * 70,
        f"  PROTOCOL QUALITY REPORT  |  v{VERSION}  |  {timestamp}",
        f"  Git commit : {commit}",
        f"  PDFs scored: {len(scored)}  |  Errors: {len(errored)}",
        "=" * 70,
        "",
    ]

    # Per-PDF sections
    for r in sorted(scored, key=lambda x: x["stem"]):
        stem  = r["stem"]
        total = r["total"]
        s1    = r["structural"]["score"]
        s2    = r["content"]["score"]
        s3    = r["formatting"]["score"]
        s4    = r["hallucination"]["score"]
        d1    = r["structural"]["detail"]
        d2    = r["content"]["detail"]
        d3    = r["formatting"]["detail"]
        d4    = r["hallucination"]["detail"]

        lines += [
            f"{'─' * 70}",
            f"  {stem}",
            f"  TOTAL: {total:5.1f} / 100   {_bar(total, 100)}",
            f"    Structural Completeness : {s1:5.1f} / 30"
            f"   ({d1['sections_found']}/{d1['sections_expected']} sections)",
            f"    Content Fidelity        : {s2:5.1f} / 40"
            f"   ({d2['values_found']}/{d2['values_total']} critical values)",
            f"    Formatting Correctness  : {s3:5.1f} / 20",
            f"    Hallucination Flag      : {s4:5.1f} / 10",
        ]

        # Missing sections
        if d1.get("missing"):
            lines.append("    Missing sections:")
            for sec in d1["missing"]:
                lines.append(f"      - {sec}")

        # Failed critical values
        if d2.get("failed"):
            lines.append("    Failed critical values:")
            for cv in d2["failed"]:
                lines.append(f"      - [{cv['id']}] {cv['description']} (expected: {cv['value']})")

        # Formatting detail
        for check, info in d3.items():
            status = "PASS" if info["passed"] else "FAIL"
            lines.append(f"    Formatting [{status}] {check}")

        # Hallucination triggers
        if d4.get("triggered"):
            lines.append("    Hallucination triggers:")
            for term in d4["triggered"]:
                lines.append(f"      ! \"{term}\"")

        lines.append("")

    # Errored files
    if errored:
        lines += ["─" * 70, "  FILES WITH ERRORS", ""]
        for r in errored:
            lines.append(f"  {r['stem']}: {r['error']}")
        lines.append("")

    # Score drift section
    history = []
    if SCORE_HISTORY_PATH.exists():
        try:
            history = json.loads(SCORE_HISTORY_PATH.read_text(encoding="utf-8"))
        except Exception:
            history = []

    if len(history) >= 2:
        prev_run = history[-2]
        prev_scores = prev_run.get("scores", {})
        prev_ts = prev_run.get("timestamp", "unknown")
        drift_lines = []
        for r in scored:
            stem = r["stem"]
            curr = r["total"]
            if stem in prev_scores:
                delta = curr - prev_scores[stem]
                if abs(delta) >= 5:
                    direction = "UP" if delta > 0 else "DOWN"
                    drift_lines.append(f"  {direction:4s} {stem:<45} {prev_scores[stem]:.1f} -> {curr:.1f}  ({delta:+.1f})")
        lines += ["=" * 70, f"  SCORE DRIFT  (vs run {prev_ts}  |  changes >= 5 pts)", "=" * 70]
        if drift_lines:
            for dl in drift_lines:
                lines.append(dl)
        else:
            lines.append("  No significant drift — all scores within 5 points of previous run.")
        lines.append("")
    else:
        lines += ["=" * 70, "  SCORE DRIFT", "=" * 70, "  Not enough history for comparison (need at least 2 runs).", ""]

    # Summary table
    lines += [
        "=" * 70,
        "  SUMMARY  (ranked by score)",
        "=" * 70,
        f"  {'PDF':<45} {'Score':>7}   Bar",
        f"  {'─'*45} {'─'*7}   {'─'*20}",
    ]
    for r in sorted(scored, key=lambda x: -x["total"]):
        lines.append(
            f"  {r['stem']:<45} {r['total']:>6.1f}   {_bar(r['total'], 100)}"
        )
    if scored:
        avg = round(sum(r["total"] for r in scored) / len(scored), 1)
        lines += [f"  {'─'*45} {'─'*7}", f"  {'Average':<45} {avg:>6.1f}"]
    lines.append("")

    # Regression risk
    at_risk = [r for r in scored if r["total"] < 60]
    lines += ["=" * 70, "  REGRESSION RISK  (score < 60)", "=" * 70]
    if at_risk:
        for r in at_risk:
            lines.append(f"  !! {r['stem']}  —  {r['total']:.1f}/100")
    else:
        lines.append("  None — all PDFs scored 60 or above.")
    lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Score history
# ---------------------------------------------------------------------------

def append_score_history(results: list[dict], timestamp: str, commit: str) -> None:
    history = []
    if SCORE_HISTORY_PATH.exists():
        try:
            history = json.loads(SCORE_HISTORY_PATH.read_text(encoding="utf-8"))
        except Exception:
            history = []

    entry = {
        "timestamp": timestamp,
        "commit":    commit,
        "scores":    {
            r["stem"]: r["total"]
            for r in results
            if "error" not in r
        },
    }
    history.append(entry)
    SCORE_HISTORY_PATH.write_text(
        json.dumps(history, indent=2, ensure_ascii=False), encoding="utf-8"
    )


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Score reference DOCX files against ground truth JSONs."
    )
    parser.add_argument(
        "--pdf",
        metavar="STEM",
        help="PDF stem (no extension) to score a single file.",
    )
    args = parser.parse_args()

    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    commit    = get_git_commit()

    if args.pdf:
        stems = [args.pdf]
    else:
        # Discover from reference_docx folder
        stems = sorted(p.stem for p in DOCX_DIR.glob("*.docx"))
        if not stems:
            sys.exit(f"[ERROR] No DOCX files found in {DOCX_DIR}")

    print(f"Scoring {len(stems)} PDF(s)  |  commit {commit}\n")

    results = []
    for stem in stems:
        print(f"  Scoring: {stem} ...", end=" ", flush=True)
        r = score_pdf(stem)
        results.append(r)
        if "error" in r:
            print(f"ERROR — {r['error']}")
        else:
            print(f"{r['total']:.1f}/100")

    # Write report
    report_text = build_report(results, timestamp, commit)
    report_path = REPORT_DIR / f"{timestamp}_report.txt"
    report_path.write_text(report_text, encoding="utf-8")
    print(f"\nReport written to: {report_path}")

    # Append score history
    append_score_history(results, timestamp, commit)
    print(f"Score history updated: {SCORE_HISTORY_PATH}")

    # Open report in default viewer
    try:
        os.startfile(str(report_path))
    except Exception as exc:
        print(f"[WARN] Could not open report automatically: {exc}")


if __name__ == "__main__":
    main()
