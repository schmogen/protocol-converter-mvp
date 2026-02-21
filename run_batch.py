import argparse
import base64
import io
import json
import logging
import re
import time
import traceback
from pathlib import Path

logger = logging.getLogger(__name__)

import pdfplumber
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from openai import APIConnectionError, APITimeoutError, InternalServerError, OpenAI, RateLimitError

# -------------------------
# Paths
# -------------------------
INPUT_DIR = Path("input_pdfs")
OUTPUT_DIR = Path("output")
CONTRACT_PATH = Path("contract.md")

# -------------------------
# Settings
# -------------------------
MODEL_CONVERT = "gpt-4.1-mini"
MODEL_FLAG = "gpt-4.1-mini"
MAX_INPUT_CHARS = 120_000  # safety limit to avoid huge uploads by accident
MAX_RETRIES = 5
RETRY_BASE_DELAY = 2  # seconds; doubled each attempt

# Always append this (deterministic), regardless of what the model outputs
REVIEW_CHECKLIST_MD = """
## Review Checklist
- [ ] Growth conditions verified (temperature, CO2/O2, time)
- [ ] Media composition and volumes verified
- [ ] Mixing parameters verified (rpm/RCF, time)
- [ ] Incubation times verified
- [ ] Centrifugation settings verified (RCF, time, temp)
- [ ] Step order confirmed against source
- [ ] Any ambiguous values flagged with [CHECK]
""".strip()

# Explicitly load .env (important on Windows)
load_dotenv(dotenv_path=".env")


# -------------------------
# PDF extraction
# -------------------------
def _clean_cell(val) -> str:
    """Normalize a pdfplumber cell value for markdown output.

    Strips leading/trailing whitespace, collapses internal newlines to a
    single space, and removes null bytes and non-printable control characters
    that can appear when pdfplumber splits cell text across column boundaries.
    """
    if val is None:
        return ""
    s = str(val)
    s = re.sub(r"[\r\n]+", " ", s)                          # newlines → space
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", s)  # non-printable
    return s.strip()


def _table_to_markdown(table_data: list) -> str:
    """Convert pdfplumber table data (list of rows) to markdown pipe format."""
    if not table_data:
        return ""
    rows = []
    for row in table_data:
        cells = [_clean_cell(cell) for cell in row]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 1:
        num_cols = len(table_data[0])
        separator = "| " + " | ".join(["---"] * num_cols) + " |"
        rows.insert(1, separator)
    return "\n".join(rows)


def _is_corrupted_table(rows: list) -> bool:
    """Return True if the table rows appear to be corrupted.

    Two corruption signals are checked against the cleaned header row:
    1. The first header cell is longer than 25 characters — indicates it has
       absorbed extra text from an adjacent body-text column.
    2. Any header cell is ≤3 characters, ends with a letter, and the next
       header cell starts with a lowercase letter — indicates a word was split
       across cells by an incorrect column boundary.
    """
    if not rows or not rows[0]:
        return False
    header = [_clean_cell(c) for c in rows[0]]
    # Signal 1: first cell is suspiciously long
    if len(header[0]) > 25:
        return True
    # Signal 2: word split across adjacent cells (e.g. "D" | "PBS volume...")
    for i in range(len(header) - 1):
        cell = header[i]
        next_cell = header[i + 1]
        if (1 <= len(cell) <= 3
                and cell[-1].isalpha()
                and next_cell
                and next_cell[0].islower()):
            return True
    # Signal 3: any header cell ends or starts with an isolated single letter
    # (e.g. "Cell Factory system D" or "T Cell Factory system p"), indicating
    # a stray character was absorbed from an adjacent body-text column.
    for cell in header:
        if re.search(r" [A-Za-z]$", cell):   # single letter tacked on at end
            return True
        if re.match(r"^[A-Za-z] ", cell):    # single letter prepended at start
            return True
    return False


def _fallback_table_rows(page, bbox: tuple):
    """Extract table rows from a bounding box using plain text extraction.

    Crops the page to bbox, extracts text, splits into non-empty lines, and
    splits each line by two or more consecutive spaces to get cells.  Returns
    a list of 2-element rows only if every line produces exactly 2 cells;
    returns None otherwise so the caller can skip the table.
    """
    x0, top, x1, bottom = bbox
    crop = page.crop((x0, top, x1, bottom))
    text = crop.extract_text() or ""
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return None
    rows = [re.split(r"  +", ln.strip()) for ln in lines]
    if not all(len(r) == 2 for r in rows):
        return None
    return rows


def _extract_tables_via_vision(page, client: OpenAI) -> list:
    """Extract tables from a page image using GPT-4o-mini vision.

    Renders the page as a JPEG at 150 DPI, sends it to the OpenAI chat
    completions API, and parses the markdown table response into a list of
    row-lists (one list per table found on the page).  Returns an empty list
    on any rendering or API error.

    Requires the 'pillow' and 'pikepdf' packages for page.to_image().
    """
    logger.info("Vision fallback triggered for page %s", getattr(page, 'page_number', '?'))
    try:
        try:
            pil_image = page.to_image(resolution=150).original
            buf = io.BytesIO()
            pil_image.save(buf, format="JPEG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
        except Exception:
            logger.error("Vision fallback: page render/encode failed:\n%s", traceback.format_exc())
            return []

        prompt = (
            "This page contains one or more tables. Extract each table you see and "
            "return them as markdown pipe-formatted tables, one after another, separated "
            "by a blank line. Include the header row and a separator row of dashes. "
            "Return only the markdown tables, nothing else."
        )

        try:
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{b64}",
                                    "detail": "high",
                                },
                            },
                            {
                                "type": "text",
                                "text": prompt,
                            },
                        ],
                    }
                ],
            )
        except Exception:
            logger.error("Vision fallback: API call failed:\n%s", traceback.format_exc())
            return []

        response_text = (resp.choices[0].message.content or "").strip()
        if not response_text:
            logger.error("Vision fallback: API returned empty response")
            return []

        # Split on blank lines; treat each chunk containing a pipe as a table
        all_tables = []
        for chunk in re.split(r"\n\s*\n", response_text):
            if "|" not in chunk:
                continue
            rows = []
            for line in chunk.splitlines():
                line = line.strip()
                if not line or not line.startswith("|"):
                    continue
                # Skip separator rows (| --- | --- |)
                if re.match(r"^[\s\-|:]+$", line.strip("|")):
                    continue
                cells = [c.strip() for c in line.split("|")[1:-1]]
                rows.append(cells)
            if rows:
                all_tables.append(rows)

        logger.info("Vision fallback: extracted %d table(s) from page %s", len(all_tables), getattr(page, 'page_number', '?'))
        return all_tables

    except Exception:
        logger.error("Vision fallback: unexpected error:\n%s", traceback.format_exc())
        return []


def _match_vision_result(corrupted_rows: list, vision_results: list, used_indices: set):
    """Find the best matching vision result for a corrupted table.

    Compares the second cell of the first row of *corrupted_rows* against the
    second cell of the first row of each unused vision result.  Normalizes both
    to lowercase with punctuation stripped, then scores by counting shared words
    of 3+ characters.  Returns the index of the best-scoring unused result with
    score > 0, or None if no match is found.
    """
    def _header_text(rows):
        if not rows or not rows[0] or len(rows[0]) < 2:
            return ""
        return _clean_cell(rows[0][1])

    def _normalize_header(text: str) -> str:
        text = text.lower()
        text = re.sub(r"[^\w\s]", "", text)
        return re.sub(r"\s+", " ", text).strip()

    def _score(a: str, b: str) -> int:
        words_a = {w for w in a.split() if len(w) >= 3}
        words_b = {w for w in b.split() if len(w) >= 3}
        return len(words_a & words_b)

    corrupted_header = _normalize_header(_header_text(corrupted_rows))

    best_idx = None
    best_score = 0
    for i, vt_rows in enumerate(vision_results):
        if i in used_indices:
            continue
        vision_header = _normalize_header(_header_text(vt_rows))
        score = _score(corrupted_header, vision_header)
        if score > best_score:
            best_score = score
            best_idx = i

    return best_idx if best_score > 0 else None


def _sort_tables_reading_order(tables: list) -> list:
    """Sort pdfplumber table objects into reading order using x-overlap detection.

    Two tables are in the same horizontal row when their x-ranges do not
    overlap (one ends at or before the other starts).  Tables are grouped
    greedily by processing them top-to-bottom: a new row begins whenever the
    new table's x-range overlaps the current row's collective x-range (i.e.
    new_x0 < row_x1 AND row_x0 < new_x1).  Non-overlapping tables (side by
    side) are collected into the same row.  Within each row tables are ordered
    left-to-right (by x0); rows are ordered top-to-bottom (by the minimum top
    of all tables in the row).
    """
    if not tables:
        return []

    # First pass: sort by top so we process tables in vertical order
    by_top = sorted(tables, key=lambda t: t.bbox[1])

    rows: list[list] = []
    current_row: list = []
    row_x0: float = 0.0
    row_x1: float = 0.0

    for table in by_top:
        x0, top, x1, bottom = table.bbox
        if current_row and (x0 < row_x1 and row_x0 < x1):
            # x-ranges overlap → this table is stacked vertically; start a new row
            rows.append(current_row)
            current_row = [table]
            row_x0 = x0
            row_x1 = x1
        elif not current_row:
            current_row = [table]
            row_x0 = x0
            row_x1 = x1
        else:
            current_row.append(table)
            row_x0 = min(row_x0, x0)
            row_x1 = max(row_x1, x1)

    if current_row:
        rows.append(current_row)

    # Sort each row left-to-right, then flatten in top-to-bottom row order
    result = []
    for row in rows:
        result.extend(sorted(row, key=lambda t: t.bbox[0]))
    return result


def _extract_page_content(page, client: OpenAI) -> str:
    """Extract text from a page using pdfplumber's plain text extraction.

    Table content is included naturally by pdfplumber's layout analysis.
    The model reads values from the extracted text and embeds them inline.
    """
    return page.extract_text(x_tolerance=3, y_tolerance=3) or ""


def _extract_page_content_with_placeholders(page, counter_start: int) -> tuple:
    """Like _extract_page_content but replaces each table with a TABLE_PLACEHOLDER_N
    token.  Returns (page_content_str, number_of_tables_found)."""
    tables = page.find_tables()
    if not tables:
        return page.extract_text() or "", 0

    tables_sorted = sorted(tables, key=lambda t: t.bbox[1])
    blocks: list[tuple[float, float, str]] = []
    counter = counter_start

    for table in tables_sorted:
        x0, top, x1, bottom = table.bbox
        rows = table.extract()
        if rows:
            blocks.append((top, bottom, f"TABLE_PLACEHOLDER_{counter}"))
            counter += 1

    n_tables = counter - counter_start

    crop_regions: list[tuple[float, float]] = []
    prev_bottom: float = 0.0
    for table in tables_sorted:
        top = table.bbox[1]
        bottom = table.bbox[3]
        if top > prev_bottom:
            crop_regions.append((prev_bottom, top))
        prev_bottom = bottom
    if prev_bottom < page.height:
        crop_regions.append((prev_bottom, page.height))

    for region_top, region_bottom in crop_regions:
        cropped = page.crop((0, region_top, page.width, region_bottom))
        text = cropped.extract_text() or ""
        if text.strip():
            blocks.append((region_top, region_bottom, text))

    blocks.sort(key=lambda b: b[0])
    return "\n\n".join(content for _, _, content in blocks), n_tables


def _looks_like_heading(line: str) -> bool:
    """Return True if the line looks like a section heading in PDF source text.

    Criteria: ≤80 chars, contains at least one letter, does not end with
    sentence punctuation, is not a bullet, and is not a numbered step.
    """
    s = line.strip()
    if not s or len(s) > 80:
        return False
    if s[0] in "-*•◗":
        return False
    if re.match(r"^\d+[.)]\s", s):
        return False
    if s[-1] in ".,:;!?":
        return False
    if not re.search(r"[a-zA-Z]", s):
        return False
    return True


def extract_tables_from_pdf(pdf_path: Path, client: OpenAI) -> list:
    """Return all tables from the PDF in page/vertical order as raw row data.

    Each entry is a dict with:
      'page'           – 1-based page number
      'top'            – top y-coordinate of the table on its page
      'rows'           – list[list[str|None]] from pdfplumber table.extract()
      'preceding_text' – text on the same page immediately above the table
      'section_heading'– most recent heading-like line seen before this table
                         across the full document, used for section matching
    """
    tables: list[dict] = []
    current_heading = ""

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            tables_on_page = _sort_tables_reading_order(page.find_tables())

            if not tables_on_page:
                # No tables on this page — scan full text to advance heading tracker
                for line in (page.extract_text() or "").splitlines():
                    if _looks_like_heading(line):
                        current_heading = line.strip()
                continue

            prev_bottom = 0.0
            _vision_results = None  # lazily populated; shared across corrupted tables on page
            _used_vision_indices: set = set()
            for table in tables_on_page:
                top = table.bbox[1]
                preceding = ""
                if top > prev_bottom:
                    crop = page.crop((0, prev_bottom, page.width, top))
                    preceding = crop.extract_text() or ""
                # Advance heading tracker from text above this table
                for line in preceding.splitlines():
                    if _looks_like_heading(line):
                        current_heading = line.strip()
                rows = table.extract()
                if rows:
                    if _is_corrupted_table(rows):
                        if _vision_results is None:
                            _vision_results = _extract_tables_via_vision(page, client)
                        match_idx = _match_vision_result(rows, _vision_results, _used_vision_indices)
                        if match_idx is not None:
                            _used_vision_indices.add(match_idx)
                            vt_rows = _vision_results[match_idx]
                            if vt_rows:
                                tables.append({
                                    "page": page_num,
                                    "top": top,
                                    "rows": vt_rows,
                                    "preceding_text": preceding.strip(),
                                    "section_heading": current_heading,
                                })
                        # else: no matching vision result; skip this corrupted table
                    else:
                        tables.append({
                            "page": page_num,
                            "top": top,
                            "rows": rows,
                            "preceding_text": preceding.strip(),
                            "section_heading": current_heading,
                        })
                prev_bottom = table.bbox[3]

            # Scan text after the last table on this page to keep tracker current
            if prev_bottom < page.height:
                remaining = page.crop((0, prev_bottom, page.width, page.height))
                for line in (remaining.extract_text() or "").splitlines():
                    if _looks_like_heading(line):
                        current_heading = line.strip()

    return tables


def extract_text_from_pdf(pdf_path: Path, client: OpenAI) -> str:
    """Extract page text with tables interleaved by vertical position."""
    parts: list[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            content = _extract_page_content(page, client)
            parts.append(f"\n\n--- PAGE {i} ---\n\n{content}")
    return "".join(parts).strip()


# -------------------------
# Cleaning helpers
# -------------------------
def remove_hyphen_linebreaks(text: str) -> str:
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text)


def normalize_newlines(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n\s*\n+", "\n<BLANKLINE>\n", text)
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    return text.replace("\n<BLANKLINE>\n", "\n\n")


def tidy_spaces(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = "\n".join(line.strip() for line in text.splitlines())
    return text.strip()


def clean_text(raw: str) -> str:
    cleaned = raw
    cleaned = remove_hyphen_linebreaks(cleaned)
    cleaned = normalize_newlines(cleaned)
    cleaned = tidy_spaces(cleaned)
    return cleaned


# Matches a TABLE_PLACEHOLDER_N token that occupies its own line.
_PLACEHOLDER_LINE_RE = re.compile(r"^(TABLE_PLACEHOLDER_\d+)$", re.MULTILINE)


def clean_text_preserving_placeholders(raw: str) -> str:
    """Like clean_text but keeps TABLE_PLACEHOLDER_N tokens on their own lines.

    clean_text's normalize_newlines step collapses single newlines to spaces,
    which would inline placeholder tokens with surrounding text.  This function
    splits the raw string on placeholder lines, applies clean_text to each
    non-placeholder segment independently, then reassembles with each placeholder
    on its own blank-line-separated paragraph.
    """
    # re.split with a capturing group interleaves text and placeholder segments.
    segments = _PLACEHOLDER_LINE_RE.split(raw)
    result_parts: list[str] = []
    for seg in segments:
        if _PLACEHOLDER_LINE_RE.match(seg.strip()):
            result_parts.append("\n\n" + seg.strip() + "\n\n")
        else:
            result_parts.append(clean_text(seg))
    return "".join(result_parts).strip()

# -------------------------
# Table reinsertion
# -------------------------
def reinsert_tables_into_markdown(md: str, pdf_tables: list) -> str:
    """Reinsert pdfplumber tables into model-generated markdown using
    section-level structural matching.

    Strategy:
    1. Group tables by their section_heading (set by extract_tables_from_pdf).
    2. Scan the model output for markdown headings and fuzzy-match each source
       section heading against them.
    3. For each matched section, insert all its tables at the end of that
       section (immediately before the next heading of equal or higher level).
    4. Tables with no section match are appended at the end of the document.
    5. TABLE_PENDING markers left by the model are stripped.

    Fuzzy heading match: normalize both to lowercase with punctuation stripped,
    then accept if one string contains the other or if they share ≥4 consecutive
    words.
    """
    # Strip TABLE_PENDING markers first so they don't affect line indexing
    lines = [ln for ln in md.split("\n") if ln.strip() != "TABLE_PENDING"]

    if not pdf_tables:
        return "\n".join(lines)

    # --- inner helpers ---

    def _heading_level(line: str) -> int:
        m = re.match(r"^(#{1,4})\s", line.strip())
        return len(m.group(1)) if m else 0

    def _heading_text(line: str) -> str:
        return re.sub(r"^#{1,4}\s+", "", line.strip())

    def _normalize(text: str) -> str:
        text = text.lower()
        text = re.sub(r"[^\w\s]", "", text)
        return re.sub(r"\s+", " ", text).strip()

    def _headings_match(src: str, md_h: str) -> bool:
        n_src, n_md = _normalize(src), _normalize(md_h)
        if not n_src or not n_md:
            return False
        if n_src in n_md or n_md in n_src:
            return True
        words = n_src.split()
        for i in range(len(words) - 3):
            if " ".join(words[i : i + 4]) in n_md:
                return True
        return False

    # --- group pdf tables by section_heading (preserve document order) ---
    section_map: dict[str, list] = {}
    unmatched: list = []
    for td in pdf_tables:
        heading = td.get("section_heading", "").strip()
        if heading:
            section_map.setdefault(heading, []).append(td)
        else:
            unmatched.append(td)

    # --- collect all markdown headings from model output ---
    md_headings: list[tuple[int, int, str]] = []  # (line_idx, level, text)
    for i, line in enumerate(lines):
        lvl = _heading_level(line)
        if lvl:
            md_headings.append((i, lvl, _heading_text(line)))

    # --- match each source section to a model heading; build insertion plan ---
    # insertion_plan: list of (insert_before_line_idx, [table_md_str, ...])
    insertion_plan: list[tuple[int, list[str]]] = []
    claimed: set[int] = set()  # model heading line indices already claimed

    for src_heading, tables_for_section in section_map.items():
        # Find first unclaimed model heading that fuzzy-matches this source heading
        match_idx = None
        match_level = None
        for h_idx, h_lvl, h_text in md_headings:
            if h_idx in claimed:
                continue
            if _headings_match(src_heading, h_text):
                match_idx = h_idx
                match_level = h_lvl
                break

        if match_idx is None:
            unmatched.extend(tables_for_section)
            continue

        claimed.add(match_idx)

        # Section ends at the next heading of equal or higher hierarchy level
        section_end = len(lines)
        for h_idx, h_lvl, _ in md_headings:
            if h_idx > match_idx and h_lvl <= match_level:
                section_end = h_idx
                break

        # Collect non-empty table markdown blocks in document order
        table_mds = [
            _table_to_markdown(td["rows"])
            for td in tables_for_section
            if _table_to_markdown(td["rows"])
        ]
        if table_mds:
            insertion_plan.append((section_end, table_mds))

    # --- apply insertions from bottom to top to avoid index drift ---
    insertion_plan.sort(key=lambda x: x[0], reverse=True)
    for insert_at, table_mds in insertion_plan:
        insert_lines: list[str] = []
        for tmd in table_mds:
            insert_lines.append("")
            insert_lines.extend(tmd.splitlines())
            insert_lines.append("")
        lines = lines[:insert_at] + insert_lines + lines[insert_at:]

    # --- append tables that had no section match at the end ---
    for td in unmatched:
        tmd = _table_to_markdown(td["rows"])
        if tmd:
            lines.append("")
            lines.extend(tmd.splitlines())
            lines.append("")

    return "\n".join(lines)


# -------------------------
# User-facing label normalization
# -------------------------
def normalize_user_facing_labels(md: str) -> str:
    """
    Normalize user-facing section titles to avoid AI-internal language.
    """
    replacements = {
        "Possible Hallucinations/Unsupported Claims": "Possible Unsupported Claims",
        "Possible Hallucinations / Unsupported Claims": "Possible Unsupported Claims",
        "Possible Hallucinations and Unsupported Claims": "Possible Unsupported Claims",
        "Possible Hallucinations": "Possible Unsupported Claims",
    }
    for old, new in replacements.items():
        md = md.replace(old, new)
    return md

# -------------------------
# Retry helper
# -------------------------
_RETRYABLE = (RateLimitError, APITimeoutError, APIConnectionError, InternalServerError)


def _call_with_retry(client: OpenAI, model: str, prompt: str) -> str:
    """Call the OpenAI responses API with exponential backoff on transient errors."""
    delay = RETRY_BASE_DELAY
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = client.responses.create(model=model, input=prompt)
            return resp.output_text.strip()
        except _RETRYABLE as e:
            if attempt == MAX_RETRIES:
                raise
            print(f"  Retry {attempt}/{MAX_RETRIES} after {type(e).__name__}, waiting {delay}s...")
            time.sleep(delay)
            delay *= 2


# -------------------------
# LLM: conversion
# -------------------------
def convert_to_protocol_markdown(client: OpenAI, contract: str, cleaned_text: str) -> str:
    if len(cleaned_text) > MAX_INPUT_CHARS:
        cleaned_text = cleaned_text[:MAX_INPUT_CHARS] + "\n\n[TRUNCATED: input exceeded MAX_INPUT_CHARS]\n"

    prompt = f"""
You are a careful scientific editor.

Follow the protocol output contract STRICTLY.
Do not add or invent scientific content.
Preserve all numbers/units/times/temperatures exactly.
If something is unclear, write "[CHECK]" rather than guessing.

STEP COMPLETENESS: Every numbered step must begin with a verb or subject and be a fully self-contained sentence. If the source text for a step begins mid-sentence, reconstruct the full sentence before outputting it. Never output a step that begins with a lowercase letter, a conjunction (such as and, or, but, so, then, while), or a word that implies a preceding clause.

SECTION HEADINGS: Output every section heading from the source document as its own standalone markdown heading using the exact wording from the source. This includes passage labels such as "Passage 1", "Passage 2", "Passage 3-5" and any other named or numbered subsections — each must appear as a separate heading and must not be merged with adjacent section titles or wrapped in parentheses as part of another heading. Do not reorder or rename any section. Do not add any section headings that do not exist in the source document, including headings such as "Objective", "Procedure", or "Materials" unless those exact words appear as headings in the source.

SOURCE FIDELITY: Do not add any sections, content, checklists, flags, or commentary that does not exist in the source document. The output must contain only what is present in the source. Do not add review checklists, parameter summaries, missing-parameter analysis, unsupported-claims sections, or any other AI-generated content.

REAGENT DEFINITIONS: When a protocol step references a named reagent mixture (e.g. "recommended growth medium", "complete medium", or any reagent given a proper name), include the full composition of that mixture either inline in that step as a sub-bullet or in a Materials section at the top of the document. Do not reference a named mixture without defining it somewhere in the document.

NUMBERED STEPS: Output every procedural step as a markdown numbered list item using the format "1. text", "2. text", and so on. Every procedural step in the source must appear as a numbered list item in the output. Do not let any step fall through to plain text.

LIST NUMBERING: Whenever a new section or subsection heading appears, any numbered list that follows must restart at 1. Each procedural section is independent and must have its own numbering starting from 1.

TABLES: Do not reproduce any tables in the output. Instead, read the data from each table in the source and embed the relevant values directly into the protocol step they belong to. For example, if a table lists wash volumes by system type, incorporate those volumes into the wash step as inline text: 'Wash with DPBS (2-layer: 80 mL, 3-layer: 120 mL, 10-layer: 400 mL, 13-layer: 520 mL per system)'. Every value from every table must appear somewhere in the output — do not omit table data.

TABLE PLACEMENT: Embed table values into the step that logically precedes or introduces that table in the source document. If a table appears after step 5, its values belong in step 5 or as a sub-note immediately after step 5.

--- CONTRACT ---
{contract}

--- INPUT TEXT ---
{cleaned_text}
""".strip()

    return _call_with_retry(client, MODEL_CONVERT, prompt)


def finalize_protocol_md(protocol_md: str) -> str:
    if "## Review Checklist" in protocol_md:
        return protocol_md.strip() + "\n"
    return protocol_md.strip() + "\n\n" + REVIEW_CHECKLIST_MD + "\n"


# -------------------------
# LLM: flagging (second pass)
# -------------------------
def generate_flags_md(client: OpenAI, cleaned_text: str, protocol_md: str) -> str:
    # Keep the inputs bounded
    if len(cleaned_text) > 60_000:
        cleaned_text = cleaned_text[:60_000] + "\n\n[TRUNCATED]\n"
    if len(protocol_md) > 60_000:
        protocol_md = protocol_md[:60_000] + "\n\n[TRUNCATED]\n"

    prompt = f"""
You are a scientific QA reviewer. Your job is to flag risk, not to rewrite.

Given:
1) SOURCE TEXT (from a PDF extraction)
2) GENERATED PROTOCOL (Markdown)

Return a Markdown report with these sections:

## Critical Parameters Checklist (from the protocol)
List the key numeric parameters present in the generated protocol (e.g., temperatures, times, volumes, concentrations, rpm/RCF, CO2/O2).
If a category is not present, say "MISSING".

## Potential Missing Parameters (compare to source)
Identify parameters that appear in the SOURCE TEXT but are missing or unclear in the GENERATED PROTOCOL.
Only cite items that are supported by the SOURCE TEXT.

## Possible Unsupported Claims
List any statements in the GENERATED PROTOCOL that are NOT clearly supported by the SOURCE TEXT.
If none, say "None detected".

## Step Order / Omission Risks
Call out any suspected omitted steps or ordering changes.
If none, say "None detected".

Rules:
- Be conservative.
- If unsure, label as "[CHECK]" instead of asserting.
- Use short bullet points.

--- SOURCE TEXT ---
{cleaned_text}

--- GENERATED PROTOCOL ---
{protocol_md}
""".strip()

    return _call_with_retry(client, MODEL_FLAG, prompt)


def append_flags_summary(protocol_md: str, flags_md: str) -> str:
    """
    Add a short flags section at the end of protocol.md so users see it immediately.
    Keep it concise by extracting just the Potential Missing + Hallucinations headers if present.
    """
    # Normalize both inputs first
    protocol_md_norm = normalize_user_facing_labels(protocol_md)
    flags_md_norm = normalize_user_facing_labels(flags_md)

    if "## Review Flags" in protocol_md_norm:
        return protocol_md_norm

    result_md = protocol_md_norm.strip() + "\n\n## Review Flags\n\n" + flags_md_norm.strip() + "\n"
    # Ensure normalization to the final concatenated string as well
    result_md = normalize_user_facing_labels(result_md)
    return result_md


# -------------------------
# Markdown -> Word export
# -------------------------
_HEADING_MAP = {"# ": 1, "## ": 2, "### ": 3}


def _parse_markdown_table(table_lines: list) -> list:
    """Parse markdown table lines into a list of rows (each a list of cell strings).
    Separator rows (| --- | --- |) are skipped."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # Skip separator row — cells contain only dashes, colons, spaces
        inner = line.strip("|")
        if re.match(r"^[\s\-|:]+$", inner):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
    return rows


def _add_table_to_doc(doc, table_lines: list) -> None:
    """Create a python-docx Table from markdown table lines."""
    rows = _parse_markdown_table(table_lines)
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    tbl = doc.add_table(rows=len(rows), cols=num_cols)
    tbl.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx in range(num_cols):
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = row.cells[c_idx]
            cell.text = cell_text
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _split_table_blocks(table_lines: list) -> list:
    """Split a flat list of consecutive pipe-delimited lines into individual
    table blocks.  A separator row (cells containing only dashes/colons) that
    appears after the first two lines of the current block signals that the
    preceding line is the header of a new table.  That line is moved into a
    fresh block together with the separator, and the old block is closed.
    """
    if not table_lines:
        return []

    def _is_sep(line: str) -> bool:
        inner = line.strip().strip("|")
        return bool(inner and re.match(r"^[\s\-|:]+$", inner))

    blocks = []
    current = []
    for line in table_lines:
        if _is_sep(line) and len(current) >= 2:
            blocks.append(current[:-1])
            current = [current[-1], line]
        else:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def md_to_docx(md_text: str, docx_path: Path, pdf_tables: list = None) -> None:
    """Convert a Markdown protocol string to a .docx file with basic formatting.

    pdf_tables – ordered list of raw table dicts from extract_tables_from_pdf().
    When present, TABLE_PLACEHOLDER_N lines are replaced with the corresponding
    pre-extracted table rendered directly from pdfplumber row data.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blank lines (Word handles paragraph spacing)
        if not stripped:
            i += 1
            continue

        # TABLE_PLACEHOLDER_N — substitute the pre-extracted PDF table
        m_ph = re.match(r"^TABLE_PLACEHOLDER_(\d+)$", stripped)
        if m_ph:
            idx = int(m_ph.group(1)) - 1  # convert 1-based N to 0-based index
            if pdf_tables and 0 <= idx < len(pdf_tables):
                md_lines = _table_to_markdown(pdf_tables[idx]["rows"]).splitlines()
                _add_table_to_doc(doc, md_lines)
            i += 1
            continue

        # Detect start of a markdown table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            for block in _split_table_blocks(table_lines):
                _add_table_to_doc(doc, block)
            continue

        # Headings
        heading_level = None
        for prefix, level in _HEADING_MAP.items():
            if stripped.startswith(prefix):
                heading_level = level
                stripped = stripped[len(prefix):]
                break
        if heading_level is not None:
            doc.add_heading(stripped, level=heading_level)
            i += 1
            continue

        # Checkbox list items: - [ ] or - [x]
        if re.match(r"^- \[[ x]\] ", stripped):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        # Bullet list items
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        # Numbered list items (e.g. "1. Step text")
        m = re.match(r"^(\d+)\.\s+", stripped)
        if m:
            doc.add_paragraph(stripped[m.end():], style="List Number")
            i += 1
            continue

        # Regular paragraph
        doc.add_paragraph(stripped)
        i += 1

    doc.save(str(docx_path))


# -------------------------
# Main
# -------------------------
def main() -> None:
    parser = argparse.ArgumentParser(description="Convert protocol PDFs to Markdown and Word.")
    parser.add_argument("--file", type=Path, help="Path to a single PDF to convert instead of the entire input folder.")
    args = parser.parse_args()

    if not CONTRACT_PATH.exists():
        raise SystemExit("Missing contract.md in project root. Create it first.")

    if args.file:
        if not args.file.exists():
            raise SystemExit(f"File not found: {args.file}")
        pdfs = [args.file]
    else:
        if not INPUT_DIR.is_dir():
            raise SystemExit(f"Input directory '{INPUT_DIR}' does not exist. Create it and add PDFs.")
        pdfs = sorted(INPUT_DIR.glob("*.pdf"))
        if not pdfs:
            raise SystemExit("No PDFs found in input_pdfs/. Add PDFs and try again.")

    contract = CONTRACT_PATH.read_text(encoding="utf-8")
    client = OpenAI()

    # Validate API key early to avoid wasting time on PDF extraction
    try:
        client.models.list()
    except Exception as e:
        raise SystemExit(f"OpenAI API key check failed: {e}") from e

    OUTPUT_DIR.mkdir(exist_ok=True)

    print(f"Found {len(pdfs)} PDFs. Batch: convert={MODEL_CONVERT}, flag={MODEL_FLAG}")

    for pdf_path in pdfs:
        t0 = time.time()
        stem = pdf_path.stem
        out_dir = OUTPUT_DIR / stem
        out_dir.mkdir(parents=True, exist_ok=True)

        log = {
            "pdf": pdf_path.name,
            "started_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "model_convert": MODEL_CONVERT,
            "model_flag": MODEL_FLAG,
            "status": "started",
        }

        try:
            raw = extract_text_from_pdf(pdf_path, client)
            (out_dir / "raw_extracted.txt").write_text(raw, encoding="utf-8")

            cleaned = clean_text(raw)
            (out_dir / "cleaned.txt").write_text(cleaned, encoding="utf-8")
            (out_dir / "cleaned_debug.txt").write_text(cleaned, encoding="utf-8")

            if len(cleaned) < 500:
                log["warning"] = "Very low extracted text. PDF may be scanned/image-only."

            protocol_md = convert_to_protocol_markdown(client, contract, cleaned)
            (out_dir / "model_output_debug.txt").write_text(protocol_md, encoding="utf-8")
            protocol_md = finalize_protocol_md(protocol_md)

            # Flagging pass (QA)
            flags_md = generate_flags_md(client, cleaned, protocol_md)
            (out_dir / "flags.md").write_text(flags_md, encoding="utf-8")

            # Append flags into protocol for visibility
            protocol_with_flags = append_flags_summary(protocol_md, flags_md)
            (out_dir / "protocol.md").write_text(protocol_with_flags, encoding="utf-8")

            # Export to Word — tables are now embedded directly in the markdown
            md_to_docx(protocol_with_flags, out_dir / "protocol.docx", [])

            log["status"] = "success"
            log["raw_chars"] = len(raw)
            log["cleaned_chars"] = len(cleaned)
            log["elapsed_seconds"] = round(time.time() - t0, 2)

            print(f"[OK] {pdf_path.name} -> {out_dir / 'protocol.md'}  ({log['elapsed_seconds']}s)")

        except Exception as e:
            log["status"] = "error"
            log["error"] = repr(e)
            log["elapsed_seconds"] = round(time.time() - t0, 2)
            print(f"[FAIL] {pdf_path.name} failed: {e!r}")

        try:
            (out_dir / "run_log.json").write_text(json.dumps(log, indent=2), encoding="utf-8")
        except OSError as e:
            print(f"[WARN] Could not write run_log.json for {pdf_path.name}: {e!r}")

    print("\nDone. Review outputs in output/<pdfname>/{protocol.md, flags.md}")


if __name__ == "__main__":
    main()
